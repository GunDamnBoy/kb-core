#!/usr/bin/env python3
"""把**已發布的**當日 JSON 機械轉成 Word 報告。

用法：podcast_docx.py <資料 repo> [日期 YYYY-MM-DD] [輸出目錄]

    podcast_docx.py ~/podcast-knowledge-digest
    podcast_docx.py ~/podcast-knowledge-digest 2026-08-20 ~/Documents/podcast-reports

## 三條刻意的設計

1. **來源是 `data/<date>.json`，不是草稿。** SKILL 第 7 步：`exit 0` 之後才產。
   這個順序是刻意的——只有真的發布成功的內容才會變成 Word，
   不會出現「Word 有這集、網站沒有」。

2. **機械轉出，不由 LLM 重寫。** 同一份內容出現在兩個地方時，
   只要有一次是「再寫一遍」，兩邊就會開始漂移，而且沒有任何訊號。
   這支程式只做排版，一個字都不改。

3. **不寫進任何 repo 目錄。** 逐字稿與 Word 檔是版權界線，而且是結構性的——
   輸出目錄預設在 `~/Documents/podcast-reports`，不是靠 `.gitignore` 擋。

## 為什麼是 python-docx 不是 docx-js

這支程式每天在 Mac 的 `~/.venvs/kb` 裡跑，跟 publish、哨兵、檢查同一個直譯器。
換成 Node 等於為了一個檔案格式在排程環境裡多養一套工具鏈，
而 `launchd` 的 PATH 是明寫的四個目錄——多一個相依就多一個會在半夜壞掉的東西。
"""
import datetime as dt
import json
import os
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

DEFAULT_OUT = "~/Documents/podcast-reports"
MUTED = RGBColor(0x6B, 0x72, 0x80)
ACCENT = RGBColor(0x1D, 0x4E, 0xD8)


def _muted(p, text, size=9):
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.font.color.rgb = MUTED
    return r


def build(doc_json: dict):
    """回傳 (文件, 實際寫進去的集數)。**兩個都要，因為第二個才是驗收依據。**"""
    d = Document()
    st = d.styles["Normal"]
    st.font.name = "Times New Roman"
    st.font.size = Pt(11)
    st.paragraph_format.space_after = Pt(6)
    st.paragraph_format.line_spacing = 1.25

    d.add_heading("Podcast 知識庫", level=0)
    p = d.add_paragraph()
    eps = doc_json.get("episodes") or []
    _muted(p, f"{doc_json.get('label', doc_json.get('date'))} · {len(eps)} 集 · "
              f"{sum(e.get('chars', 0) for e in eps):,} 字", 10)

    cc = doc_json.get("crossCut") or {}
    if cc.get("points"):
        d.add_heading(cc.get("title", "跨節目交叉觀察"), level=1)
        if cc.get("intro"):
            d.add_paragraph(cc["intro"])
        for pt in cc["points"]:
            d.add_heading(pt.get("title", ""), level=2)
            d.add_paragraph(pt.get("body", ""))

    ps = doc_json.get("postscript") or {}
    if ps.get("paragraphs"):
        d.add_heading(ps.get("title", "本期觀察後記"), level=1)
        for para in ps["paragraphs"]:
            d.add_paragraph(para)
    for o in ps.get("observations") or []:
        q = d.add_paragraph(style="Intense Quote")
        q.add_run(o.get("text", ""))
        sub = d.add_paragraph()
        _muted(sub, f"怎麼驗：{o.get('check','')}　·　到期 {o.get('horizon','')}"
                    f"　·　{o.get('status','')}")

    written = 0
    for e in eps:
        d.add_page_break()
        d.add_heading(e.get("title", "").split("｜", 1)[-1], level=1)

        meta = d.add_paragraph()
        bits = [e.get("show", ""), e.get("published", ""),
                f"{e.get('chars', 0):,} 字", "、".join(e.get("topics") or [])]
        _muted(meta, " · ".join(b for b in bits if b))
        if e.get("guest"):
            g = d.add_paragraph()
            _muted(g, f"來賓：{e['guest']}")

        d.add_paragraph(e.get("summary", ""))
        body_from = len(d.paragraphs) - 1   # 從摘要那一段起算，才是這一集的內容

        if e.get("takeaways"):
            d.add_heading("核心重點", level=2)
            for t in e["takeaways"]:
                h = d.add_paragraph()
                lab = h.add_run(f"{t.get('label','')}｜")
                lab.bold = True
                lab.font.color.rgb = ACCENT
                h.add_run(t.get("title", "")).bold = True
                d.add_paragraph(t.get("body", ""))

        if e.get("sections"):
            d.add_heading("完整摘譯", level=2)
            for s in e["sections"]:
                if s.get("heading"):
                    d.add_heading(s["heading"], level=3)
                for para in s.get("paragraphs") or []:
                    d.add_paragraph(para)

        if e.get("quotes"):
            d.add_heading("金句", level=2)
            for q in e["quotes"]:
                b = d.add_paragraph(style="Quote")
                b.add_run(q.get("text", ""))
                by = d.add_paragraph()
                by.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                _muted(by, f"— {q.get('by','')}")

        # 一集「有寫進去」的定義是**文件真的多了內容**，不是迴圈跑了一輪。
        # 原本這裡是迴圈頂端無條件 written += 1，那樣 written 恆等於 len(eps)，
        # 底下那個 written != expected 的守衛永遠不可能成立——
        # 它看起來在數集數，其實只是在數迴圈圈數。
        # 界線取「有沒有字」而不是「幾個字」：空殼是 0，任何一集正常內容都遠大於 0，
        # 中間沒有需要調的門檻。字數夠不夠是 chars_in_tier 在發布前的事，不是這裡。
        if sum(len(x.text.strip()) for x in d.paragraphs[body_from:]):
            written += 1

        qa = e.get("quality") or {}
        for label, key in (("講者", "speakerNote"), ("時間軸", "timestampNote")):
            if qa.get(key):
                n = d.add_paragraph()
                _muted(n, f"{label}：{qa[key]}", 8)
        src = d.add_paragraph()
        _muted(src, f"{e.get('source','')}　{e.get('url','')}", 8)

    tail = d.add_paragraph()
    _muted(tail, "本文為個人知識管理用途的中文摘譯，非投資建議。"
                 "每集均附原節目連結，鼓勵前往收聽。", 8)
    return d, written


def main(argv) -> int:
    if not 2 <= len(argv) <= 4:
        print(__doc__)
        return 2
    repo = Path(argv[1]).expanduser()
    date = argv[2] if len(argv) >= 3 else dt.datetime.now(
        dt.timezone(dt.timedelta(hours=8))).strftime("%Y-%m-%d")
    outdir = Path(argv[3]).expanduser() if len(argv) == 4 else Path(
        os.path.expanduser(DEFAULT_OUT))

    src = repo / "data" / f"{date}.json"
    if not src.exists():
        # 「還沒發布」與「發布了但轉檔失敗」是兩件事，訊息要分得出來。
        print(f"找不到 {src} —— 那一天還沒發布，Word 不該先於網站存在",
              file=sys.stderr)
        return 13

    # **輸出目錄不能在任何 repo 裡。** 這是版權界線，不靠 .gitignore。
    if (outdir / ".git").exists() or any(
            (p / ".git").exists() for p in outdir.parents if p != p.parent):
        print(f"拒絕寫入 {outdir} —— 它在一個 git repo 裡面", file=sys.stderr)
        return 2

    outdir.mkdir(parents=True, exist_ok=True)
    out = outdir / f"podcast-{date}.docx"
    doc_json = json.loads(src.read_text())
    d, written = build(doc_json)
    d.save(out)

    expected = len(doc_json.get("episodes") or [])
    size = out.stat().st_size
    print(f"{out}　{written} 集　{size:,} bytes")

    # **集數才是驗收依據，位元組數只是第二道。**
    # 舊規格明文要求「腳本會印出集數，要等於當日實際集數」，而我第一版只看位元組——
    # 一份少了三集但每集都很長的報告，位元組數看起來完全正常。
    # 今天第三次撞到同一個形狀：守衛量的維度比它宣稱保護的東西低一階。
    if written != expected:
        print(f"寫進去 {written} 集，當日應有 {expected} 集 —— 少的那幾集去哪了",
              file=sys.stderr)
        return 10
    if size < 20_000:
        print("位元組數異常地小 —— 檔案存在不等於內容在裡面", file=sys.stderr)
        return 10
    return 0


if __name__ == "__main__":
    sys.exit(main(sys.argv))

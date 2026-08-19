#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""json2docx — 把節目知識庫的每日 JSON 轉成 Word 報告。

2026-08-08 之前，每日排程寫完 data/YYYY-MM-DD.json 後還要用 LLM 把同一份內容
重寫成 82 頁的 Word——**同一份內容付兩次輸出 token**。這支腳本把第二次變成
機械轉檔：內容以 JSON 為唯一來源，docx 只是排版。

用法：
    python3 ~/.podfetch/json2docx.py <data/YYYY-MM-DD.json> <輸出.docx>

相依：python-docx（沙箱：pip install python-docx --break-system-packages）
輸出**絕不可放進 repo**（Public），寫到暫存輸出資料夾再交付。
"""

import json
import sys


def main():
    if len(sys.argv) != 3:
        print(__doc__)
        return 1
    src, dst = sys.argv[1], sys.argv[2]

    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        print("缺 python-docx：pip install python-docx --break-system-packages")
        return 1

    d = json.load(open(src, encoding="utf-8"))
    doc = Document()

    # 基本樣式：繁中正文
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)
    # 東亞字型要另外設
    from docx.oxml.ns import qn
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang TC")

    ACCENT = RGBColor(0x1E, 0x40, 0xAF)
    MUTED = RGBColor(0x64, 0x74, 0x8B)

    def h(text, level, color=ACCENT, size=None):
        p = doc.add_heading("", level=level)
        r = p.add_run(text)
        r.font.color.rgb = color
        if size:
            r.font.size = Pt(size)
        r.font.name = "Calibri"
        r.element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang TC")
        return p

    def para(text, bold=False, color=None, size=None, italic=False):
        p = doc.add_paragraph()
        r = p.add_run(text)
        r.bold = bold
        r.italic = italic
        if color:
            r.font.color.rgb = color
        if size:
            r.font.size = Pt(size)
        return p

    # ── 封面 ──
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("節目知識庫・每日摘譯")
    r.bold = True
    r.font.size = Pt(24)
    r.font.color.rgb = ACCENT
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("%s｜共 %d 集" % (d.get("label", d.get("date", "")),
                                      len(d.get("episodes", []))))
    r.font.size = Pt(13)
    r.font.color.rgb = MUTED
    shows = "、".join(dict.fromkeys(e.get("show", "") for e in d.get("episodes", [])))
    sub2 = doc.add_paragraph()
    sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub2.add_run(shows)
    r.font.size = Pt(10)
    r.font.color.rgb = MUTED
    para("本文件為個人知識管理用途之濃縮摘譯，非投資建議。引用數字前請回查原始節目。",
         color=MUTED, size=9, italic=True)

    # ── 跨節目交叉觀察 ──
    cc = d.get("crossCut")
    if cc:
        doc.add_page_break()
        h(cc.get("title", "跨節目交叉觀察"), 1)
        if cc.get("intro"):
            para(cc["intro"])
        for pt in cc.get("points", []):
            h(pt.get("title", ""), 2, size=13)
            para(pt.get("body", ""))

    ps = d.get("postscript")
    if ps:
        h(ps.get("title", "本期觀察後記"), 1)
        for pg in ps.get("paragraphs", []):
            para(pg)

    # ── 各集 ──
    for e in d.get("episodes", []):
        doc.add_page_break()
        h(e.get("title", e.get("show", "")), 1)
        meta_bits = []
        for m in e.get("meta", []):
            meta_bits.append("%s：%s" % (m.get("k", ""), m.get("v", "")))
        # A 第 7 節要求每集保留原文連結。meta 裡沒有「原始連結」時，
        # 用頂層 url 補——連結不能取決於子代理有沒有把它放進 meta。
        if e.get("url") and not any("連結" in m.get("k", "") for m in e.get("meta", [])):
            meta_bits.append("原始連結：%s" % e["url"])
        for key, label in (("published", None), ("hosts", "主持"),
                           ("guest", "來賓"), ("source", "來源")):
            v = e.get(key)
            if v:
                meta_bits.append("%s：%s" % (label, v) if label else v)
        if meta_bits:
            para("｜".join(meta_bits), color=MUTED, size=9)

        if e.get("summary"):
            p = para("一句話總結　", bold=True)
            p.add_run(e["summary"])

        tk = e.get("takeaways", [])
        if tk:
            h("核心重點", 2, size=13)
            for t in tk:
                p = doc.add_paragraph()
                r = p.add_run("%s｜%s　" % (t.get("label", ""), t.get("title", "")))
                r.bold = True
                p.add_run(t.get("body", ""))

        secs = e.get("sections", [])
        if secs:
            h("完整摘譯", 2, size=13)
            for s in secs:
                if s.get("heading"):
                    para(s["heading"], bold=True)
                for pg in s.get("paragraphs", []):
                    para(pg)

        qs = e.get("quotes", [])
        if qs:
            h("本集金句", 2, size=13)
            for q in qs:
                p = doc.add_paragraph()
                r = p.add_run("「%s」" % q.get("text", ""))
                r.italic = True
                if q.get("by"):
                    r2 = p.add_run("　—— %s" % q["by"])
                    r2.font.color.rgb = MUTED

    doc.save(dst)
    n_ep = len(d.get("episodes", []))
    print("完成：%s（%d 集）" % (dst, n_ep))
    return 0


if __name__ == "__main__":
    sys.exit(main())
